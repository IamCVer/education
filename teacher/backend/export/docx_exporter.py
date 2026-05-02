from __future__ import annotations

from pathlib import Path
from typing import Dict, List


def export_lesson_plan_docx(output_path: Path, sections: Dict[str, List[str] | str]) -> Path:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx 未安装，无法导出 docx") from exc

    document = Document()
    title = sections.get("教学主题", "AI 教学设计")
    document.add_heading(str(title), level=0)
    for heading, content in sections.items():
        if heading == "教学主题":
            continue
        document.add_heading(heading, level=1)
        if isinstance(content, list):
            for item in content:
                document.add_paragraph(str(item), style="List Bullet")
        else:
            document.add_paragraph(str(content))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path
